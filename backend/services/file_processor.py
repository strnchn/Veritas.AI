"""
Serviço de processamento de arquivos (PDF, DOCX, TXT)
"""
import os
from typing import Tuple
from pathlib import Path
import PyPDF2
from docx import Document


class FileProcessor:
    """Classe para processar diferentes tipos de arquivos"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Obtém a extensão do arquivo"""
        return Path(filename).suffix.lower()
    
    @staticmethod
    def is_allowed_file(filename: str) -> bool:
        """Verifica se o arquivo tem extensão permitida"""
        return FileProcessor.get_file_extension(filename) in FileProcessor.ALLOWED_EXTENSIONS
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extrai texto de arquivo PDF usando PyPDF2
        """
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extrai texto de todas as páginas
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n\n"
                    except Exception as e:
                        print(f"Aviso: Erro ao extrair página {page_num + 1}: {e}")
                        continue
            
            if not text.strip():
                raise Exception("Não foi possível extrair texto do PDF. O arquivo pode estar vazio ou protegido.")
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extrai texto de arquivo DOCX"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extrai texto dos parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extrai texto das tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n\n".join(text)
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extrai texto de arquivo TXT"""
        try:
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Não foi possível decodificar o arquivo TXT")
        except Exception as e:
            raise Exception(f"Erro ao ler arquivo TXT: {str(e)}")
    
    @classmethod
    def process_file(cls, file_path: str, filename: str) -> Tuple[str, str]:
        """
        Processa arquivo e retorna o texto extraído
        
        Args:
            file_path: Caminho do arquivo
            filename: Nome original do arquivo
            
        Returns:
            Tuple[str, str]: (texto_extraído, tipo_arquivo)
        """
        if not cls.is_allowed_file(filename):
            raise ValueError(f"Tipo de arquivo não permitido. Extensões permitidas: {cls.ALLOWED_EXTENSIONS}")
        
        extension = cls.get_file_extension(filename)
        
        try:
            if extension == '.pdf':
                text = cls.extract_text_from_pdf(file_path)
                file_type = "PDF"
            elif extension == '.docx':
                text = cls.extract_text_from_docx(file_path)
                file_type = "DOCX"
            elif extension == '.txt':
                text = cls.extract_text_from_txt(file_path)
                file_type = "TXT"
            else:
                raise ValueError(f"Extensão não suportada: {extension}")
            
            if not text or len(text.strip()) < 100:
                raise ValueError("O arquivo está vazio ou contém muito pouco texto para análise")
            
            return text, file_type
        
        except Exception as e:
            raise Exception(f"Erro ao processar arquivo {filename}: {str(e)}")

